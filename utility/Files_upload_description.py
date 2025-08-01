import os
from openai import OpenAI
from google.cloud import storage
import PyPDF2
from io import BytesIO
import logging
from itertools import chain
from langchain.text_splitter import CharacterTextSplitter
from dotenv import load_dotenv  # Import the dotenv module
import pymongo
import json
from bson import ObjectId
from docx import Document as DocxDocument
from langchain_core.documents import Document
import pandas as pd

# Setup logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# MongoDB setup
mongo_client = pymongo.MongoClient("mongodb://dev:N47309HxFWE2Ehc@35.209.224.122:27017")
db = mongo_client["ChatbotDB-DEV"]
collection = db['files']

# Load environment variables from .env file
load_dotenv()  # This will load the .env file in your project directory

# Retrieve OpenAI API Key from environment variables
openai_api_key = os.getenv("OPENAI_API_KEY")  # Use the key stored in .env

# Check if the API key is loaded correctly
if not openai_api_key:
    raise ValueError("OpenAI API key is not set. Please set it in the .env file.")

# Initialize OpenAI client
client = OpenAI(api_key=openai_api_key)

# Initialize the text splitter
text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

# Function to generate description, keywords, and tags using OpenAI API
def generate_openai_output(text):
    prompt = f"""Please read the following text or pdf, excel, doc, etc  data  and write a description , keywords and tags as  given below provide the response in JSON format:
    {{
        "description": "10-15 word description of the content",
        "keywords": ["5 important keywords"],
        "tags": ["relevant tags such as product name, purpose"]
    }}

    {text}"""

    # Generate the output using OpenAI's GPT-4 model
    response = client.chat.completions.create(
        model="gpt-4",  # Use GPT-4 model
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,  # Lower temperature for more deterministic output
    )
    
    # Correctly handle the OpenAI response
    content = response.choices[0].message.content

    
    
    
    
    # Extract the description, keywords, and tags from the response
   
    print (content)
    return content

# Function to read PDFs from GCS and extract text
# MongoDB setup
mongo_client = pymongo.MongoClient("mongodb://dev:N47309HxFWE2Ehc@35.209.224.122:27017")
db = mongo_client["ChatbotDB-DEV"]
collection = db['files']

def read_pdf_from_gcs(bucket_name, blob_names, chatbot_id, version_id):
    """Read PDFs from GCS and extract text with error handling"""
    complete_document = []
    print(f"Processing the following blobs: {blob_names}")
    
    # Initialize the storage client once
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    
    for blob_name in blob_names:
        try:
            logger.info(f"Processing blob: {blob_name}")
            blob = bucket.blob(blob_name)
            pdf_bytes = blob.download_as_bytes()
            pdf_file = BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            data = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    data.append(text)
                else:
                    logger.warning(f"No text extracted from page {page_num} in blob {blob_name}")
            
            pdf_text = '. '.join(data)

            # Generate Description, Keywords, and Tags using OpenAIs
            file_des = generate_openai_output(pdf_text)

            file_des = json.loads(file_des)

            print(file_des)

            description = file_des.get('description', 'No description')
            keywords = file_des.get('keywords', [])
            tags = file_des.get('tags', [])

            # Print the generated description, keywords, and tags
            print(f"Generated description for {blob_name}: {description}")
            print(f"Generated keywords for {blob_name}: {keywords}")
            print(f"Generated tags for {blob_name}: {tags}")

            logger.info(f"Generated description: {description}")
            logger.info(f"Generated keywords: {keywords}")
            logger.info(f"Generated tags: {tags}")

            chatbot_oid = ObjectId(chatbot_id)
            version_oid = ObjectId(version_id)

            # Update MongoDB collection with description, keywords, and tags
            collection.update_one(
                {"chatbot_id": chatbot_oid, "version_id": version_oid, "file_name" : blob_name},
                {
                    "$set": {
                        "description": description,
                        "keywords": keywords,
                        "tags": tags
                    }
                },
                upsert=True  # Create a new document if no match is found
            )
            logger.info(f"MongoDB updated for blob: {blob_name}")

        except Exception as e:
            logger.error(f"Error processing blob {blob_name}: {e}")
    
    # Return the processed documents
    flattened_docs = list(chain.from_iterable(complete_document))
    return flattened_docs

def document_splitter(text: str) :
    # Assuming you have a text splitting logic, which uses SemanticChunker
    try:
        long_doc = [Document(page_content=text)]
        # Use your document splitter here
        docs = text_splitter.split_documents(long_doc)
        return docs
    except Exception as e:
        logger.error(f"Document creator failed: {str(e)}")
        raise

def read_documents_from_gcs(bucket_name, blob_names, chatbot_id, version_id):
    """Read various document types from GCS, extract text, enrich with OpenAI metadata, and update MongoDB"""
    try:
        all_docs = []
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)

        for blob_name in blob_names:
            logger.info(f"📂 Processing blob: {blob_name}")
            blob = bucket.blob(blob_name)

            if not blob.exists():
                logger.warning(f"⚠️ Blob '{blob_name}' not found in bucket '{bucket_name}'. Skipping.")
                continue

            try:
                file_bytes = blob.download_as_bytes()
                ext = os.path.splitext(blob_name)[-1].lower()
                logger.info(f"📄 Detected file type: {ext}")

                # Text extraction based on file type
                if ext == '.pdf':
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                    text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

                elif ext == '.txt':
                    text = file_bytes.decode('utf-8', errors='ignore')

                elif ext == '.docx':
                    doc = DocxDocument(BytesIO(file_bytes))
                    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                elif ext in ['.xls', '.xlsx']:
                    excel = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
                    text = "\n\n".join([
                        f"Sheet: {sheet}\n{df.to_string(index=False)}"
                        for sheet, df in excel.items()
                    ])

                else:
                    logger.warning(f"❌ Unsupported file format: {ext} for '{blob_name}'. Skipping.")
                    continue

                # Split document into smaller chunks
                docs = document_splitter(text)
                all_docs.append(docs)

                # Generate metadata from OpenAI
                openai_response = generate_openai_output(docs)
                try:
                    file_meta = json.loads(openai_response)
                    description = file_meta.get('description', 'No description')
                    keywords = file_meta.get('keywords', [])
                    tags = file_meta.get('tags', [])
                except Exception as json_err:
                    logger.warning(f"⚠️ Failed to parse OpenAI response: {openai_response}")
                    description, keywords, tags = 'No description', [], []

                logger.info(f"📝 Description for {blob_name}: {description}")
                logger.info(f"🔑 Keywords: {keywords}")
                logger.info(f"🏷️ Tags: {tags}")

                # Update MongoDB
                collection.update_one(
                    {
                        "chatbot_id": ObjectId(chatbot_id),
                        "version_id": ObjectId(version_id),
                        "file_name": blob_name
                    },
                    {
                        "$set": {
                            "description": description,
                            "keywords": keywords,
                            "tags": tags
                        }
                    },
                    upsert=True
                )
                logger.info(f"✅ MongoDB updated for blob: {blob_name}")

            except Exception as file_err:
                logger.error(f"❌ Error processing blob {blob_name}: {str(file_err)}")
                continue

        return list(chain.from_iterable(all_docs))

    except Exception as e:
        logger.error(f"🔥 Document reading pipeline failed: {str(e)}")
        raise


# Function to process each file content
def description_from_gcs(bucket_name, blob_names,chatbot_id,version_id):
    """Load PDF documents from GCS and return them as documents"""
    try:
        docs = read_documents_from_gcs(bucket_name, blob_names,chatbot_id,version_id)
        if not docs:
            logger.warning("No documents were extracted from the PDFs.")
            return "No documents extracted."

        # Return the documents
        return docs

    except Exception as e:
        logger.error(f"Failed to process GCS files: {e}")
        return f"Error processing GCS files: {e}"

# Main function to run the process
 # Print the extracted documents


